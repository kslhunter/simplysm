"""DOCX handler: extract text, images, tables, and OLE embedded objects."""

from _common import ensure_packages, ext_from_content_type, normalize_cell, parse_heading_level

PACKAGES = {"python-docx": "docx"}


def extract(file_path):
    ensure_packages(PACKAGES)
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    text_parts = []
    images = []
    embedded = []

    for para in doc.paragraphs:
        for run in para.runs:
            drawings = (run._element.findall(f".//{qn('wp:inline')}") +
                        run._element.findall(f".//{qn('wp:anchor')}"))
            for drawing in drawings:
                blip = drawing.find(f".//{qn('a:blip')}")
                if blip is not None:
                    embed_id = blip.get(qn("r:embed"))
                    if embed_id:
                        rel = doc.part.rels.get(embed_id)
                        if rel and hasattr(rel, 'target_part'):
                            ext = ext_from_content_type(rel.target_part.content_type)
                            images.append({
                                "data": rel.target_part.blob,
                                "ext": ext,
                                "context": "paragraph image",
                            })

        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            prefix = ""
            if "Heading" in style:
                level = parse_heading_level(style)
                if level is not None:
                    prefix = "#" * level + " "
                else:
                    prefix = "## "
            text_parts.append(f"{prefix}{text}")

    for t_idx, table in enumerate(doc.tables):
        text_parts.append(f"\n### Table {t_idx + 1}\n")
        for row in table.rows:
            cells = [normalize_cell(cell.text) for cell in row.cells]
            text_parts.append("| " + " | ".join(cells) + " |")

    # OLE embedded objects
    seen = set()
    for rel in doc.part.rels.values():
        reltype = rel.reltype or ""
        if "oleObject" in reltype or "package" in reltype:
            target_ref = getattr(rel, 'target_ref', '') or ''
            if target_ref in seen:
                continue
            seen.add(target_ref)
            try:
                blob = rel.target_part.blob
                filename = target_ref.split("/")[-1] if "/" in target_ref else target_ref
                if not filename:
                    filename = f"embedded_{len(embedded) + 1}.bin"
                embedded.append({"filename": filename, "data": blob})
            except Exception:
                pass

    return {
        "text": "\n".join(text_parts),
        "images": images,
        "embedded": embedded,
        "metadata": {},
    }
