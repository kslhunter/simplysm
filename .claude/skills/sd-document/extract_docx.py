#!/usr/bin/env python3
"""Extract text and images from DOCX files in paragraph flow order."""

from _common import (
    setup_encoding, make_output_paths, print_header, save_image,
    ext_from_content_type, print_image_summary, run_cli,
    normalize_cell, parse_heading_level,
)

setup_encoding()


def extract(file_path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    fp, out_dir = make_output_paths(file_path)
    img_idx = 0

    print_header(fp)

    for para in doc.paragraphs:
        for run in para.runs:
            drawings = run._element.findall(f".//{qn('wp:inline')}") + run._element.findall(f".//{qn('wp:anchor')}")
            for drawing in drawings:
                blip = drawing.find(f".//{qn('a:blip')}")
                if blip is not None:
                    embed = blip.get(qn("r:embed"))
                    if embed:
                        rel = doc.part.rels.get(embed)
                        if rel:
                            img_idx += 1
                            ext = ext_from_content_type(rel.target_part.content_type)
                            img_path = save_image(out_dir, img_idx, rel.target_part.blob, ext)
                            print(f"[IMG] {img_path}")

        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            prefix = ""
            if "Heading" in style:
                level = parse_heading_level(style)
                if level is not None:
                    prefix = "#" * level + " "
                else:
                    prefix = "## "
            print(f"{prefix}{text}")

    # Table extraction
    for t_idx, table in enumerate(doc.tables):
        print(f"\n### Table {t_idx + 1}\n")
        for row in table.rows:
            cells = [normalize_cell(cell.text) for cell in row.cells]
            print("| " + " | ".join(cells) + " |")

    print()
    print_image_summary(img_idx, out_dir)


if __name__ == "__main__":
    run_cli(extract, "extract_docx.py", {"python-docx": "docx"})
