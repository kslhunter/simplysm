"""DOCX handler: extract text, images, tables, and OLE embedded objects."""

from _common import ensure_packages, ext_from_content_type, normalize_cell, parse_heading_level

PACKAGES = {"python-docx": "docx"}


def extract(file_path):
    ensure_packages(PACKAGES)
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph

    doc = Document(file_path)
    text_parts = []
    images = []
    embedded = []
    img_idx = 0
    emb_idx = 0

    def _extract_drawing(drawing):
        nonlocal img_idx
        blip = drawing.find(f".//{qn('a:blip')}")
        if blip is None:
            return None
        embed_id = blip.get(qn("r:embed"))
        if not embed_id:
            return None
        rel = doc.part.rels.get(embed_id)
        if not rel or not hasattr(rel, 'target_part'):
            return None
        ext = ext_from_content_type(rel.target_part.content_type)
        img_idx += 1
        doc_pr = drawing.find(f".//{qn('wp:docPr')}")
        alt = ""
        if doc_pr is not None:
            alt = doc_pr.get("descr", "") or doc_pr.get("title", "")
        images.append({
            "data": rel.target_part.blob,
            "ext": ext,
            "context": alt or "paragraph image",
        })
        return img_idx

    def _process_paragraph(element):
        para = Paragraph(element, doc)
        style = para.style.name if para.style else ""
        prefix = ""
        if "Heading" in style:
            level = parse_heading_level(style)
            prefix = "#" * (level or 2) + " "

        parts = []
        for run in para.runs:
            if run.text:
                parts.append(run.text)
            drawings = (run._element.findall(f".//{qn('wp:inline')}") +
                        run._element.findall(f".//{qn('wp:anchor')}"))
            for d in drawings:
                idx = _extract_drawing(d)
                if idx is not None:
                    parts.append(f"[IMG:{idx}]")

        line = "".join(parts).strip()
        if line:
            text_parts.append(f"{prefix}{line}")

    def _process_table(element):
        table = DocxTable(element, doc)
        rows = list(table.rows)
        if not rows:
            return
        text_parts.append("")
        for r_idx, row in enumerate(rows):
            cells = [normalize_cell(cell.text) for cell in row.cells]
            text_parts.append("| " + " | ".join(cells) + " |")
            if r_idx == 0:
                text_parts.append("|" + "|".join(["---"] * len(cells)) + "|")
        text_parts.append("")

    # Iterate body elements in document order (paragraphs and tables interleaved)
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            _process_paragraph(child)
        elif tag == 'tbl':
            _process_table(child)

    # Headers and footers
    for sec_idx, section in enumerate(doc.sections):
        h_parts = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
        f_parts = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
        if h_parts or f_parts:
            text_parts.append("")
            text_parts.append(f"[Header/Footer — Section {sec_idx + 1}]")
            if h_parts:
                text_parts.append(f"Header: {' | '.join(h_parts)}")
            if f_parts:
                text_parts.append(f"Footer: {' | '.join(f_parts)}")

    # OLE embedded objects
    seen = set()
    for rel in doc.part.rels.values():
        reltype = rel.reltype or ""
        if "oleObject" in reltype or "package" in reltype:
            target_ref = getattr(rel, 'target_ref', '') or ''
            if target_ref in seen:
                continue
            seen.add(target_ref)
            try:
                blob = rel.target_part.blob
                filename = target_ref.split("/")[-1] if "/" in target_ref else target_ref
                if not filename:
                    filename = f"embedded_{len(embedded) + 1}.bin"
                emb_idx += 1
                embedded.append({"filename": filename, "data": blob})
                text_parts.append(f"[EMB:{emb_idx}]")
            except Exception:
                pass

    return {
        "text": "\n".join(text_parts),
        "images": images,
        "embedded": embedded,
        "metadata": {},
    }
